from pathlib import Path
from docx import Document
from .base import BaseParser, register
from app.models import ParsedDocument, DocumentMetadata
import hashlib


@register(".docx", ".doc")
class DocxParser(BaseParser):
    """Parser for Microsoft Word DOCX files using python-docx."""

    def parse(self, path: Path) -> ParsedDocument:
        """Extract text and metadata from a DOCX file."""
        document = Document(path)

        # Extract text from paragraphs
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        text = "\n\n".join(paragraphs)

        # Add text from tables
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
                if row_text:
                    text += f"\n{row_text}"

        # Generate content hash ID
        content_hash = hashlib.md5(text.encode()).hexdigest()

        # Extract document properties if available
        core_props = document.core_properties

        # Create document metadata
        document_metadata = DocumentMetadata(
            filename=path.name,
            size_bytes=path.stat().st_size,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            title=core_props.title if hasattr(core_props, "title") and core_props.title else None,
            author=core_props.author
            if hasattr(core_props, "author") and core_props.author
            else None,
            created=core_props.created
            if hasattr(core_props, "created") and core_props.created
            else None,
            modified=core_props.modified if hasattr(core_props, "modified") else None,
            paragraphs=len(document.paragraphs),
            tables=len(document.tables),
        )

        return ParsedDocument(
            id=content_hash,
            filename=path.name,
            path=str(path.absolute()),
            text=text,
            metadata=document_metadata,
        )
